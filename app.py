import streamlit as st
from docx import Document
from ai_paper_generator import generate_multiple

st.set_page_config(
    page_title="AI Question Paper Generator",
    page_icon="🎓",
    layout="wide"
)

# ------------------ UI ------------------

st.markdown("""
# 🎓 AI Question Paper Generator
Generate AI-created university question papers using TensorFlow
""")

st.markdown("---")

col1, col2, col3 = st.columns(3)

with col1:
    college = st.text_input(
        "College Name",
        "CIT College"
    )

with col2:
    department = st.text_input(
        "Department",
        "Computer Science"
    )

with col3:
    semester = st.selectbox(
        "Semester",
        [1,2,3,4,5,6,7,8]
    )

exam_name = st.text_input(
    "Exam Name",
    "Model Examination"
)

st.markdown("---")

subjects = [
    "AI",
    "ML",
    "CN",
    "OS",
    "SE"
]

blueprint = {}

st.subheader("Custom Blueprint")

for subject in subjects:

    st.markdown(f"### {subject}")

    c1, c2 = st.columns(2)

    with c1:
        two_marks = st.number_input(
            f"{subject} - 2 Marks",
            min_value=0,
            max_value=20,
            value=0,
            key=f"{subject}_2"
        )

    with c2:
        ten_marks = st.number_input(
            f"{subject} - 10 Marks",
            min_value=0,
            max_value=20,
            value=0,
            key=f"{subject}_10"
        )

    blueprint[subject] = {
        "2": two_marks,
        "10": ten_marks
    }

# ------------------ GENERATE ------------------

if st.button("🚀 Generate AI Question Paper"):

    paper = []

    for subject in subjects:

        two_count = blueprint[subject]["2"]
        ten_count = blueprint[subject]["10"]

        # 2 MARK QUESTIONS

        if two_count > 0:

            questions = generate_multiple(
                subject,
                2,
                two_count
            )

            for q in questions:

                paper.append({
                    "Question": q,
                    "Marks": 2,
                    "Subject": subject
                })

        # 10 MARK QUESTIONS

        if ten_count > 0:

            questions = generate_multiple(
                subject,
                10,
                ten_count
            )

            for q in questions:

                paper.append({
                    "Question": q,
                    "Marks": 10,
                    "Subject": subject
                })

    total_marks = sum(
        q["Marks"]
        for q in paper
    )

    st.success(
        f"Question Paper Generated Successfully | Total Marks = {total_marks}"
    )

    st.markdown(f"""
## 🏫 {college}

### Department of {department}

### {exam_name}

### Semester {semester}

# TOTAL MARKS : {total_marks}
""")

    # ---------------- PART A ----------------

    st.markdown("## PART A (2 Marks)")

    qno = 1

    for q in paper:

        if q["Marks"] == 2:

            st.write(
                f"{qno}. {q['Question']}"
            )

            qno += 1

    st.markdown("---")

    # ---------------- PART B ----------------

    st.markdown("## PART B (10 Marks)")

    for q in paper:

        if q["Marks"] == 10:

            st.write(
                f"{qno}. {q['Question']}"
            )

            qno += 1

    st.markdown("---")

    st.subheader(
        f"TOTAL MARKS : {total_marks}"
    )

    # ---------------- DOCX ----------------

    doc = Document()

    doc.add_heading(
        college,
        level=1
    )

    doc.add_paragraph(
        f"Department : {department}"
    )

    doc.add_paragraph(
        f"Exam : {exam_name}"
    )

    doc.add_paragraph(
        f"Semester : {semester}"
    )

    doc.add_paragraph(
        f"Total Marks : {total_marks}"
    )

    doc.add_heading(
        "PART A (2 Marks)",
        level=2
    )

    qno = 1

    for q in paper:

        if q["Marks"] == 2:

            doc.add_paragraph(
                f"{qno}. {q['Question']}"
            )

            qno += 1

    doc.add_heading(
        "PART B (10 Marks)",
        level=2
    )

    for q in paper:

        if q["Marks"] == 10:

            doc.add_paragraph(
                f"{qno}. {q['Question']}"
            )

            qno += 1

    filename = "AI_Question_Paper.docx"

    doc.save(filename)

    with open(
        filename,
        "rb"
    ) as file:

        st.download_button(
            "📥 Download DOCX",
            file,
            file_name=filename
        )

st.markdown("---")

st.caption(
    "Built using TensorFlow, NLP and Streamlit"
)