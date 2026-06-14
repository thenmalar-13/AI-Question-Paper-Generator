from docx import Document

doc = Document()

doc.add_heading("QUESTION PAPER", 0)

doc.add_paragraph("PART A")

doc.add_paragraph("1. What is AI?")

doc.add_paragraph("2. Define Machine Learning")

doc.add_paragraph("")

doc.add_paragraph("PART B")

doc.add_paragraph("3. Explain Neural Networks")

doc.add_paragraph("4. Describe Deadlock Handling")

doc.save("question_paper.docx")

print("DOCX Created")